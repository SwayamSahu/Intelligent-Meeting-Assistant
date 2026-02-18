#pip install moviepy python-docx
import os, time
import pandas as pd
from moviepy import VideoFileClip
from docx import Document
from docx.shared import Inches

# ────────────────────────────────
# Helpers
# ────────────────────────────────
def parse_timestamp(ts):
    """
    Accepts either
      • int / float              (already seconds), or
      • "HH:MM:SS[.mmm]" / "MM:SS[.mmm]" strings.
    Returns: float seconds.
    """
    if isinstance(ts, (int, float)):
        return float(ts)

    parts = str(ts).strip().split(":")
    if len(parts) == 3:              # HH:MM:SS(.mmm)
        h, m, s = parts
    elif len(parts) == 2:            # MM:SS(.mmm)  → treat HH = 0
        h = 0
        m, s = parts
    else:
        raise ValueError(f"Bad timestamp: {ts!r}")

    return float(h) * 3600 + float(m) * 60 + float(s)


def seconds_to_hhmmss(sec):
    """Turn *sec* into 'HH-MM-SS' (dashes are filename-safe)."""
    whole, frac = divmod(sec, 1)
    base = time.strftime("%H-%M-%S", time.gmtime(whole))
    return f"{base}.{int(frac*1000):03d}" if frac else base

# # ────────────────────────────────
# # Auto timestamp detection
# # ────────────────────────────────
# def get_auto_timestamps(csv_path, video_path, min_gap_sec=2):
#     """
#     Reads event CSV and returns a list of timestamps (seconds) 
#     for interesting events (clicks, ROI changes, etc.),
#     ensuring at least `min_gap_sec` between them.
#     """
#     if not os.path.exists(csv_path):
#         print(f"⚠ CSV file not found: {csv_path}")
#         return []

#     df = pd.read_csv(csv_path)

#     # Adjust this filter to match actual CSV column names and event types
#     interesting_events = df[df["event"].str.upper() == "CLICK"]

#     if interesting_events.empty:
#         print("⚠ No CLICK events found in CSV.")
#         return []

#     # Get FPS from the video
#     clip = VideoFileClip(video_path)
#     fps = clip.fps
#     clip.close()

#     timestamps = []
#     last_time = -min_gap_sec

#     for frame_name in interesting_events["frame"]:
#         sec = int(frame_name.replace(".jpg", "")) / fps
#         if sec - last_time >= min_gap_sec:
#             timestamps.append(sec)
#             last_time = sec

#     print(f"Auto-selected {len(timestamps)} CLICK timestamps for GIF report: {timestamps}")
#     return timestamps

# ────────────────────────────────
# Extract common timestamps
# ────────────────────────────────
def get_common_timestamps(xlsx_path, csv_path):
    """
    Reads:
      - Excel → timestamp_start
      - CSV → frame_timestamps where event == CLICK
    Returns: list of common timestamps (in seconds).
    """
    if not os.path.exists(xlsx_path) or not os.path.exists(csv_path):
        print("⚠ Missing input files")
        return []

    # Read Excel timestamps
    df_xlsx = pd.read_excel(xlsx_path)
    xlsx_times = [parse_timestamp(t) for t in df_xlsx["timestamp_start"].dropna()]

    # Read CSV timestamps (CLICK only)
    df_csv = pd.read_csv(csv_path)
    click_times = [parse_timestamp(t) for t in df_csv.loc[df_csv["event"].str.upper() == "CLICK", "frame_timestamps"]]

    # Find intersection (within 1 sec tolerance for safety)
    common = []
    for xt in xlsx_times:
        for ct in click_times:
            if abs(xt - ct) < 1:   # tolerance 1 sec
                common.append(ct)

    print(f"Found {len(common)} common timestamps: {common}")
    return common


# ────────────────────────────────
# Core functions 
# ────────────────────────────────
def extract_gif_snip(video_path, center_sec, out_path, padding=10):
    clip = VideoFileClip(video_path)
    sub = clip.subclipped(max(center_sec - padding, 0),
                       min(center_sec + padding, clip.duration))
    sub.resized(width=720).write_gif(out_path, fps=10)
    clip.close()
    print(f"GIF saved → {out_path}")


# def insert_gifs_to_word(gif_dir, out_docx="GIF_Report.docx"):
#     doc = Document()
#     doc.add_heading("GIF Snippets Report", level=1)

#     for f in sorted(p for p in os.listdir(gif_dir) if p.lower().endswith(".gif")):
#         doc.add_heading(os.path.splitext(f)[0], level=2)
#         doc.add_picture(os.path.join(gif_dir, f), width=Inches(3))
#         doc.add_paragraph()

#     doc.save(out_docx)
#     print(f"DOCX saved → {out_docx}")


def insert_gifs_to_word(gif_dir, transcript_csv, out_docx="GIF_Report.docx"):
    """
    Insert gifs into DOCX with transcript text (from transcript_csv) as title above each gif.
    """
    # Load transcript CSV
    df_transcript = pd.read_csv(transcript_csv)
    df_transcript["timestamp_sec"] = df_transcript["timestamp_start"].apply(parse_timestamp)

    doc = Document()
    doc.add_heading("GIF Snippets Report", level=1)

    for f in sorted(p for p in os.listdir(gif_dir) if p.lower().endswith(".gif")):
        # Extract timestamp from filename
        ts_str = os.path.splitext(f)[0].replace("-", ":").split(".")[0]
        try:
            gif_sec = parse_timestamp(ts_str.replace(":", "-").replace("-", ":"))
        except:
            gif_sec = None

        # Find nearest transcript text
        caption = ""
        if gif_sec is not None:
            nearest_row = df_transcript.iloc[(df_transcript["timestamp_sec"] - gif_sec).abs().argsort()[:1]]
            caption = nearest_row["text"].values[0] if not nearest_row.empty else ""

        # Insert into Word
        if caption:
            doc.add_heading(caption, level=2)

        # Insert timestamp as a normal paragraph
        doc.add_paragraph(ts_str)
        # Insert GIF below    
        doc.add_picture(os.path.join(gif_dir, f), width=Inches(3))
        doc.add_paragraph()

    doc.save(out_docx)
    print(f"DOCX saved → {out_docx}")


# ────────────────────────────────
# Orchestrator
# ────────────────────────────────
def create_gif_report(video_path, xlsx_path, csv_path, transcript_csv, out_dir="out/gifs", out_docx="out/GIF_Report.docx", padding=10):
    """
    Creates GIFs for timestamps common between:
      - output.xlsx (timestamp_start)
      - events.csv (CLICK events)
    Adds transcript text above each GIF in the DOCX report.  
    """
    timestamps = get_common_timestamps(xlsx_path, csv_path)

    if not timestamps:
        print("⚠ No common timestamps found. Skipping GIF report.")
        return

    os.makedirs(out_dir, exist_ok=True)

    for ts in timestamps:
        sec = parse_timestamp(ts)
        fname = seconds_to_hhmmss(sec) + ".gif"
        extract_gif_snip(video_path, sec, os.path.join(out_dir, fname), padding)

    insert_gifs_to_word(out_dir, transcript_csv, out_docx)


# ────────────────────────────────
# Example usage
# ────────────────────────────────
if __name__ == "__main__":
    video_path = r"data\windows_detect.mp4"
    xlsx_path  = r"output\output.xlsx"
    csv_path   = r"out\events.csv"
    transcript_csv = r"output\audio_transcript_timestamps.csv"
    create_gif_report(video_path, xlsx_path, csv_path, transcript_csv)  # auto mode

    # OR manual mode:
    # timestamps = ["00:00:05", "00:00:15", "00:00:25"]
    # create_gif_report(video_path, timestamps=timestamps)
